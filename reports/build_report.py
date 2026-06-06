"""
Generate reports/report.docx — first draft of the Shopify Sales & Weather Analysis report.
Run from repo root: python reports/build_report.py

All figures are read from figures/.  All numbers are sourced from the analysis
notebooks and check scripts; none are invented.
"""
import sys
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT    = Path(__file__).parent.parent
FIGURES = ROOT / "figures"
REPORTS = ROOT / "reports"

# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_font(run, size_pt=11, bold=False, italic=False, colour=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if colour:
        run.font.color.rgb = RGBColor(*colour)


def _add_page_break(doc):
    doc.add_paragraph().add_run().add_break(
        __import__("docx.enum.text", fromlist=["WD_BREAK_TYPE"]).WD_BREAK_TYPE.PAGE
    )


def _heading(doc, text, level):
    """Add a heading.  level 1 = section, level 2 = subsection."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    return p


def _body(doc, text, space_after=8):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Pt(0)
    # Justify
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


fig_counter = [1]   # mutable so the closure can update it

def _figure(doc, filename, caption, width=Inches(5.8)):
    """Embed a figure with a numbered italicised caption."""
    path = FIGURES / filename
    if not path.exists():
        _body(doc, f"[FIGURE NOT FOUND: {filename}]")
        return

    # Image — centred
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.space_before = Pt(6)
    pic_para.add_run().add_picture(str(path), width=width)

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    run = cap.add_run(f"Figure {fig_counter[0]}: {caption}")
    _set_font(run, size_pt=9.5, italic=True, colour=(80, 80, 80))
    fig_counter[0] += 1


def _bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    return p


# ── Document setup ────────────────────────────────────────────────────────────

doc = Document()

# Margins: 2.5 cm all around
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

doc.core_properties.title   = "Shopify Sales & Weather Analysis"
doc.core_properties.subject = "Descriptive Sales Analysis and Weather-Driven Revenue Modelling"
doc.core_properties.author  = "Data Analysis Team"

# ── Title page ────────────────────────────────────────────────────────────────

# Vertical space
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Shopify Sales & Weather Analysis")
_set_font(title_run, size_pt=28, bold=True, colour=(31, 73, 125))

doc.add_paragraph()

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run(
    "Descriptive Sales Analysis and Weather-Driven Revenue Modelling\n"
    "for a Danish Artisan Food Chain"
)
_set_font(sub_run, size_pt=14, colour=(68, 114, 196))

for _ in range(3):
    doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(date.today().strftime("%B %Y"))
_set_font(date_run, size_pt=12, colour=(89, 89, 89))

conf_p = doc.add_paragraph()
conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
conf_run = conf_p.add_run("Confidential — not for distribution")
_set_font(conf_run, size_pt=10, italic=True, colour=(150, 150, 150))

_add_page_break(doc)

# ── Section 1: Data & Methodology ────────────────────────────────────────────

_heading(doc, "1  Data & Methodology", level=1)

_heading(doc, "1.1  Data Source", level=2)
_body(doc,
    "Transaction data were extracted from the client's Shopify platform via the "
    "Shopify Admin API.  The export covers the period from January 2024 to March 2025 "
    "(15 calendar months) and comprises seven tables: orders, line_items, products, "
    "locations, customers, discounts, and refunds.  After cleaning and joining, the "
    "analytical dataset contains 454,821 transaction lines across 282,380 unique orders "
    "placed across seven physical stores."
)
_body(doc,
    "The central fact table is line_items, which records one row per product line "
    "within each order.  It was joined to: orders (timestamps, financial totals, "
    "fulfilment status, and store location); products (product title, vendor, and tags, "
    "joined on product_id); locations (store name and city, joined on location_id); and "
    "refunds (summed per order and filled with zero where no refund was recorded).  "
    "Customers and discounts were loaded but are not used in the analyses presented here."
)

_heading(doc, "1.2  Cleaning Rules", level=2)
_body(doc,
    "The following rules were applied to the merged dataset in sequence, "
    "yielding the 454,821-row analytical base:"
)
_bullet(doc, "Columns with more than 95% missing values were dropped (sku, customer_id, email, product_type).")
_bullet(doc, "Missing categoricals were filled: store_name → 'Unknown Store'; title → 'Unknown Product'; vendor → 'Unknown Vendor'; tags → 'Unknown Tag'.")
_bullet(doc, "Rows tagged 'Indpakning' (packaging-only lines) were removed.")
_bullet(doc, "Only completed sales were retained: financial_status = 'paid' AND fulfillment_status = 'fulfilled' AND cancelled_at is null.")
_bullet(doc, "Rows from a designated test store were excluded.")
_bullet(doc, "A transaction-level revenue column was derived as price × quantity.")
_body(doc,
    "The product name field (name) from line_items is used as the primary product "
    "identifier throughout, rather than the title field from the products catalogue.  "
    "The name field is populated by the point-of-sale (POS) terminal for every "
    "transaction, whereas title depends on a successful product_id join to the "
    "Shopify catalogue — a join that fails for approximately 1% of lines (see Section 3)."
)

_heading(doc, "1.3  Product Categorisation", level=2)
_body(doc,
    "Products were assigned to one of eight client-defined families based on the "
    "POS name field, using an exact-name lookup against a curated mapping table "
    "(src/product_categories.csv).  The eight families are: Ice Cream, Hot Beverages, "
    "Buns & Bakery, Chocolate, Gifts & Cards, Snacks & Nuts, Christmas, and Others.  "
    "The mapping incorporates client feedback that moved flødeboller and bar-type "
    "products from Buns & Bakery to Chocolate, which changed the category ranking "
    "(Chocolate overtook Buns & Bakery to become the second-largest category by revenue)."
)

_heading(doc, "1.4  Anonymisation Note", level=2)
_body(doc,
    "Store identities are anonymised throughout this report.  The seven stores are "
    "referred to as Store A through Store G, ranked by total revenue (Store A = "
    "highest revenue).  Revenue figures and unit counts are expressed as a percentage "
    "of the dataset total, or as an index scaled so that the maximum value equals 100, "
    "rather than in absolute DKK.  This protects client confidentiality while "
    "preserving all relative comparisons and analytical conclusions."
)

_add_page_break(doc)

# ── Section 2: Sales Analysis ─────────────────────────────────────────────────

_heading(doc, "2  Sales Analysis", level=1)

_heading(doc, "2.1  Product Concentration", level=2)
_body(doc,
    "Revenue is highly concentrated in a small number of products.  The single "
    "top product — a two-scoop ice cream — accounts for approximately 19.6% of "
    "total revenue, more than any other product by a substantial margin.  "
    "The next largest, a one-scoop ice cream, contributes around 10.8%.  "
    "Together the top five products account for roughly half of all revenue."
)
_figure(doc, "top_15_products_by_revenue.png",
    "Top 15 products by revenue share (% of total).  "
    "The two- and one-scoop ice creams dominate; the remaining products in the top 15 "
    "each contribute 1–7%.")

_body(doc,
    "A Pareto analysis confirms the classic 80/20 pattern: 21 products collectively "
    "account for 80% of total revenue, out of 320 distinct product names in the "
    "catalogue.  The cumulative revenue curve rises steeply through the first 20 "
    "products and then flattens sharply, indicating that the long tail of niche "
    "products contributes comparatively little to overall sales."
)
_figure(doc, "pareto_revenue.png",
    "Pareto chart: product revenue index (bars, left axis) and cumulative revenue "
    "share (red line, right axis) for the top 30 products.  The dashed red line "
    "marks the 80% threshold; the vertical dotted line falls at product 21.")

_heading(doc, "2.2  Revenue by Category", level=2)
_body(doc,
    "Ice Cream is by far the dominant revenue category, contributing 57.0% of total "
    "revenue.  Chocolate is the second-largest at 17.0%, followed by Buns & Bakery at "
    "12.4%.  Together these three categories account for more than 86% of all revenue.  "
    "Hot Beverages (6.0%), Gifts & Cards (3.6%), and Others (2.4%) make up most of the "
    "remainder, while Snacks & Nuts (0.9%) and Christmas (0.8%) are small in absolute "
    "terms but may serve important margin and basket-building roles."
)
_body(doc,
    "The dominance of Ice Cream is consistent with the client's positioning as a "
    "premium artisan ice cream chain.  The relative size of Chocolate — driven largely "
    "by flødeboller (cream puffs) and chocolate-coated bars — reflects a secondary "
    "product identity that is strongest in the colder months when ice cream demand falls."
)
_figure(doc, "category_revenue_breakdown.png",
    "Revenue share by product category (% of total).  "
    "Ice Cream (57%) and Chocolate (17%) together account for nearly three-quarters "
    "of all revenue.")

_heading(doc, "2.3  Seasonality", level=2)

_body(doc,
    "Monthly revenue follows a clear seasonal arc driven by ice cream demand.  "
    "May 2024 was the peak month, with June 2024 also significantly above the annual "
    "average.  Revenue then softens through July and August before recovering partially "
    "in August — consistent with summer holiday patterns in Denmark — before declining "
    "through the autumn.  Winter months (October through February) are substantially "
    "quieter, running at roughly one-quarter to one-third of peak-month levels."
)
_figure(doc, "monthly_revenue.png",
    "Monthly revenue index (max = 100, May 2024).  The seasonal arc peaks in "
    "May–June, falls through autumn, and recovers modestly in December.")

_body(doc,
    "Day-of-week patterns show that Saturday and Friday are the two busiest trading "
    "days by revenue, generating roughly 60–70% more revenue than the quietest days "
    "(Monday and Tuesday).  This pattern is consistent across all stores and reflects "
    "the discretionary, leisure-oriented nature of the product range — customers "
    "visit artisan food stores more frequently on weekends and towards the end of the "
    "working week."
)
_figure(doc, "weekday_revenue.png",
    "Revenue by day of week (index, max = 100 = Saturday).  Friday and Saturday "
    "are the two peak days; Monday and Tuesday are the quietest.")

_body(doc,
    "Within the trading day, revenue is concentrated in the early-to-mid afternoon.  "
    "The peak revenue window is 14:00–16:00 (local Copenhagen time), with 15:00 and "
    "16:00 being the top two hours by revenue index.  Activity rises from store opening "
    "around 10:00 and is substantially lower in the morning and evening.  Staffing and "
    "preparation schedules should be calibrated to this window."
)
_figure(doc, "hourly_revenue.png",
    "Revenue by hour of day, Copenhagen local time (index, max = 100).  "
    "The 14:00–16:00 window drives the largest share of daily revenue.")

_body(doc,
    "Category seasonality reveals the structural logic behind the aggregate curve.  "
    "Ice Cream peaks sharply in May and June, contributing over 75% of revenue in those "
    "months, and falls back substantially from October onwards.  Hot Beverages and "
    "Christmas items follow an inverse seasonal profile, peaking in November and "
    "December when ice cream demand is lowest.  This counter-seasonality is commercially "
    "important: it moderates the winter revenue trough and keeps cross-trained staff "
    "productively employed year-round.  Buns & Bakery is more evenly distributed, "
    "serving as a year-round complement to both ice cream and hot drinks."
)
_figure(doc, "category_seasonality.png",
    "Category revenue heatmap: each cell shows that category's revenue as a "
    "percentage of its own annual total.  Ice Cream peaks May–June; "
    "Hot Beverages and Christmas peak November–December.",
    width=Inches(6.0))

_heading(doc, "2.4  Store Comparison", level=2)
_body(doc,
    "Store A is the highest-revenue store, accounting for approximately 21.7% of "
    "total revenue, with Store B (17.4%) and Store C (17.1%) close behind.  "
    "Store G is the smallest by both revenue (6.0%) and order count (7.3%).  "
    "Average order value (AOV) varies meaningfully across stores: Store B records the "
    "highest AOV (indexed at 100), suggesting either higher unit prices, larger basket "
    "sizes, or a stronger mix of premium products.  Store G records the lowest AOV, "
    "which may reflect a different customer profile or product mix."
)
_figure(doc, "store_revenue.png",
    "Store performance overview: revenue share (%), order share (%), and average "
    "order value index (max = 100 = Store B).  "
    "Store A leads on revenue; Store B on transaction value.",
    width=Inches(6.2))

_body(doc,
    "The category mix charts reveal that Ice Cream is the dominant category at every "
    "store, but the degree of dominance varies.  Store A has the highest Ice Cream "
    "share (consistent with it being in a high-footfall area), while stores with lower "
    "Ice Cream dependence show proportionally stronger Chocolate and Buns & Bakery "
    "contributions.  Niche categories such as Snacks & Nuts and Christmas show "
    "variation across locations, suggesting locally tailored assortments or "
    "different customer demographics."
)
_figure(doc, "store_category_mix.png",
    "Category revenue mix by store (100% stacked bar, % of each store's revenue).  "
    "Ice Cream dominates everywhere; relative Chocolate and Buns & Bakery shares "
    "vary materially by location.")

_add_page_break(doc)

# ── Section 3: Data Quality ───────────────────────────────────────────────────

_heading(doc, "3  Data Quality — Unknown Products", level=1)

_heading(doc, "3.1  Definition of the Issue", level=2)
_body(doc,
    "Each row in the analytical dataset carries two product identity fields: name, "
    "populated directly by the POS terminal at the point of sale, and title, sourced "
    "from the Shopify product catalogue via a join on product_id.  When a POS terminal "
    "records a transaction without a valid product_id — for example, when a product "
    "is sold under a temporary or unregistered SKU — the catalogue join fails, and "
    "title is filled with the placeholder value 'Unknown Product'.  "
    "The name field is always present and is used for all categorisation and analysis "
    "in this report; the Unknown Product flag therefore does not compromise revenue or "
    "category findings.  It does, however, signal a POS configuration gap that is "
    "worth investigating and correcting."
)

_heading(doc, "3.2  Scale and Row-vs-Unit Divergence", level=2)
_body(doc,
    "At the row level, Unknown Product is a minor issue: 4,484 lines out of 454,821 "
    "carry the flag, a rate of 0.99%.  At the unit (quantity-weighted) level, however, "
    "the rate is 11.77% — approximately twelve times larger.  The reason is that "
    "Unknown Product lines are not typical single-item transactions: they average "
    "approximately 28 units per line, compared with roughly 2 units per line for "
    "known-product lines.  These are bulk or box-format sales — ice cream tubs, "
    "multi-pack confectionery, or similar high-quantity items — that were entered at "
    "the POS without a catalogue-linked SKU.  The unit rate is therefore the more "
    "meaningful indicator of the operational footprint of this issue."
)
_figure(doc, "dq_unknown_by_store.png",
    "Unknown Product unit rate by store.  Store A carries 35.8% of units "
    "as Unknown Product — far above any other location.")

_heading(doc, "3.3  Store Concentration", level=2)
_body(doc,
    "The issue is heavily concentrated in one store.  Store A accounts for 92,393 "
    "of the 125,717 unknown units — approximately 73% — despite contributing 21.7% "
    "of total revenue.  Its unknown unit rate is 35.8%, compared with 10.5% for the "
    "second-most-affected store (Store D) and under 4% for all others.  "
    "This pattern strongly suggests a store-specific POS configuration problem: a "
    "terminal at Store A was accepting high-quantity transactions for products not "
    "registered in the Shopify catalogue, rather than a system-wide data collection issue."
)

_heading(doc, "3.4  Trend Over Time", level=2)
_body(doc,
    "The time series of the unknown unit rate does not show a single, clean cutover.  "
    "Instead, it is highly irregular: the overall rate peaks at 29.5% in June 2024 "
    "and at 24.3% in May 2024, falls sharply in July 2024, rises again in August 2024 "
    "(11.4%), and then declines erratically before settling near zero from late 2024 "
    "onwards.  The row rate — less sensitive to the bulk-quantity effect — shows a more "
    "gradual improvement from about 2.6% in January 2024 to under 0.3% by "
    "February 2025."
)
_figure(doc, "dq_unknown_over_time.png",
    "Unknown Product rate over time for all stores combined: unit rate (solid red) "
    "and row rate (dashed).  The unit rate is highly variable with large summer peaks; "
    "both rates converge near zero by early 2025.")

_body(doc,
    "At Store A specifically, the unit rate reaches 67.5% in June 2024 and 54.1% "
    "in May 2024, implying that more than half of all units sold at that store in "
    "those months were not linked to any Shopify catalogue entry.  The rate drops below "
    "1% from July 2024 in most months, though individual months see small spikes "
    "through to December 2024.  From January 2025 onwards the rate is consistently "
    "below 0.5% at this store."
)
_figure(doc, "dq_unknown_store_over_time.png",
    "Unknown Product rate over time for Store A (most-affected): unit rate (solid) "
    "and row rate (dashed).  Peak unit rates of 54–68% in May–June 2024 point to "
    "intermittent high-volume POS entries without catalogue linkage.")

_body(doc,
    "The pattern is consistent with a POS terminal that intermittently accepted "
    "bulk product entries — possibly from a secondary interface or offline mode — "
    "without resolving the SKU against the Shopify catalogue.  The resolution appears "
    "to have been implemented progressively rather than as a single configuration "
    "change, with the most severe episodes concentrated in the summer peak period "
    "when transaction volumes are highest.  A full investigation should verify whether "
    "any revenues were misattributed to incorrect product categories during the "
    "affected period."
)

# ── Save ──────────────────────────────────────────────────────────────────────

out_path = REPORTS / "report.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")
print(f"  Figures embedded: {fig_counter[0] - 1}")
print(f"  Pages: ~{len(doc.paragraphs) // 18 + 1} (estimated)")
