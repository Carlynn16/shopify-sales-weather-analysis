"""
Finalize report.docx:
  1. Harmonize period wording: "15 calendar months" -> "14-month period (January 2024-March 2025)"
  2. Export to report.pdf via docx2pdf (uses installed Word)
  3. Safety-check PDF text: 0 DKK, 0 real suburb/client names

Run from project root:  python scripts/finalize_report.py
"""
import re
import sys
from pathlib import Path

from docx import Document

DOCX = Path("report.docx")
PDF  = Path("report.pdf")

OLD_PHRASE = "the period from January 2024 to March 2025 (15 calendar months)"
NEW_PHRASE = "a 14-month period (January 2024–March 2025)"  # en-dash


# ---------------------------------------------------------------------------
# Step 1: Fix period wording in docx
# ---------------------------------------------------------------------------

def fix_period_wording(doc: Document) -> int:
    """Replace OLD_PHRASE with NEW_PHRASE across all runs in Normal paragraphs.
    Returns the number of replacements made.
    """
    replacements = 0
    for para in doc.paragraphs:
        if OLD_PHRASE not in para.text:
            continue
        # Rebuild the paragraph text across its runs via inline replacement.
        # python-docx stores text in runs; a phrase may span multiple runs.
        # Simplest safe approach: reassemble full text, replace, write back to
        # the first run, clear the rest.
        full = "".join(r.text for r in para.runs)
        if OLD_PHRASE not in full:
            continue
        new_full = full.replace(OLD_PHRASE, NEW_PHRASE, 1)
        # Put the corrected text into the first run and blank the others.
        if para.runs:
            para.runs[0].text = new_full
            for r in para.runs[1:]:
                r.text = ""
        replacements += 1
    return replacements


# ---------------------------------------------------------------------------
# Step 2: Export to PDF
# ---------------------------------------------------------------------------

def export_pdf() -> None:
    from docx2pdf import convert
    print(f"Converting {DOCX} -> {PDF} (opens Word in background) ...")
    convert(str(DOCX), str(PDF))
    size_kb = PDF.stat().st_size // 1024
    print(f"PDF written: {PDF}  ({size_kb:,} KB)")


# ---------------------------------------------------------------------------
# Step 3: Safety-check PDF text
# ---------------------------------------------------------------------------

def pdf_page_count() -> int:
    from pdfminer.high_level import extract_pages
    return sum(1 for _ in extract_pages(str(PDF)))


def extract_pdf_text() -> str:
    from pdfminer.high_level import extract_text
    return extract_text(str(PDF))


def safety_check(text: str) -> bool:
    suspects = {
        "DKK":           r"\bDKK\b",
        "Vedbaek/Vedb":  r"Vedb",
        "Hellerup":      r"Hellerup",
        "Gentofte":      r"Gentofte",
        "Lyngby":        r"Lyngby",
        "Charlottenlund":r"Charlottenlund",
        "Klampenborg":   r"Klampenborg",
    }
    all_clear = True
    for label, pat in suspects.items():
        hits = re.findall(pat, text, re.IGNORECASE)
        status = "FAIL" if hits else "ok"
        if hits:
            all_clear = False
        print(f"  [{status}] {label}: {len(hits)}")
    return all_clear


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    # ── 1. Wording fix ───────────────────────────────────────────────────────
    doc = Document(DOCX)
    n = fix_period_wording(doc)
    if n == 0:
        # Maybe already fixed or phrasing differs — report and continue
        print("WARNING: period-wording phrase not found; docx unchanged")
    else:
        doc.save(DOCX)
        print(f"Wording fix applied ({n} replacement(s)).")

    # Verify the fix in a fresh load
    doc2 = Document(DOCX)
    for p in doc2.paragraphs:
        if "calendar month" in p.text.lower() or "14-month" in p.text.lower() \
                or "january 2024" in p.text.lower():
            print(f"  Section 1.1 period wording: {p.text[:120]}")
            break

    # ── 2. PDF export ────────────────────────────────────────────────────────
    export_pdf()

    # ── 3. Verify PDF ────────────────────────────────────────────────────────
    pages = pdf_page_count()
    print(f"\nPDF page count: {pages}")

    print("\nExtracting PDF text for safety check ...")
    pdf_text = extract_pdf_text()
    print(f"Characters extracted: {len(pdf_text):,}")

    print("\nSafety scan (PDF):")
    ok = safety_check(pdf_text)

    # Quick figure presence check: count "Figure" references
    fig_refs = len(re.findall(r"\bFigure\s+\d+", pdf_text))
    print(f"\nFigure references found in PDF text: {fig_refs}  (expected ~20)")

    print(f"\n{'PASS' if ok else 'FAIL'} — safety check {'clear' if ok else 'has hits — review above'}.")


if __name__ == "__main__":
    main()
