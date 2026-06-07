"""
Append 'Section 5: Predictive Modelling' to reports/report.docx.
Run from project root: python scripts/append_modelling_section.py
All numbers grounded in actual modeling.py outputs; no absolute DKK.
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

REPORT  = Path("reports/report.docx")
FIGURES = Path("figures")
W       = Inches(6.3)

_HEADER_BG = "D9E1F2"
_FONT_HDR  = Pt(9)
_FONT_BODY = Pt(9)
_FONT_CAP  = Pt(10)


# ---------------------------------------------------------------------------
# Helpers (identical style to insert_tables.py)
# ---------------------------------------------------------------------------

def _set_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_color)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    tcPr.append(shd)


def _cell(cell, text: str, bold: bool = False, size: Pt = _FONT_BODY,
          align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(str(text))
    run.bold      = bold
    run.font.size = size


def _style_header(table) -> None:
    for cell in table.rows[0].cells:
        _set_bg(cell, _HEADER_BG)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold      = True
                run.font.size = _FONT_HDR


def _set_col_widths(table, widths_in) -> None:
    for row in table.rows:
        for i, w in enumerate(widths_in):
            row.cells[i].width = Inches(w)


def add_para(doc, text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_figure(doc, filename: str, caption: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURES / filename), width=W)
    cap = doc.add_paragraph(style="Normal")
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(12)
    cap.add_run(caption)


def add_table_caption(doc, text: str) -> None:
    cap = doc.add_paragraph(style="Normal")
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after  = Pt(2)
    run = cap.add_run(text)
    run.bold      = True
    run.font.size = _FONT_CAP


def add_spacer(doc) -> None:
    doc.add_paragraph(style="Normal").paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Section 5
# ---------------------------------------------------------------------------

def build_section(doc: Document) -> None:

    # ── 5  Predictive Modelling ───────────────────────────────────────────
    doc.add_heading("5  Predictive Modelling", level=1)

    add_para(doc,
        "This section describes two complementary models built to quantify and "
        "forecast daily Ice Cream revenue at the store-day level.  The first model "
        "is an interpretable log-linear regression (the 'why') that translates each "
        "predictor into an explicit percentage effect on revenue.  The second is an "
        "XGBoost gradient-boosting forecaster (the 'how much') that captures "
        "non-linear interactions and autoregressive dynamics.  Both models use the "
        "VIF-pruned feature set established in Section 4.2, plus store fixed effects "
        "and calendar controls; all metrics are scale-free (R², WAPE, percentage "
        "effects) -- no absolute revenue figures are reported."
    )

    # ── 5.1  Interpretable Regression ─────────────────────────────────────
    doc.add_heading("5.1  Interpretable Regression (OLS)", level=2)

    add_para(doc,
        "An OLS model is fitted to log(Ice Cream revenue + 1) with heteroscedasticity-"
        "consistent HC3 standard errors.  The log-linear specification was preferred "
        "over a Gamma GLM with log link for three reasons: (1) the +1 offset is "
        "negligible relative to typical revenue magnitudes, making log(revenue + 1) "
        "approximately equal to log(revenue) for all practical values; (2) only 1.7% "
        "of store-days carry zero Ice Cream revenue, so zero-inflation treatment is "
        "unnecessary; and (3) OLS on the log scale yields closed-form HC3 confidence "
        "intervals and coefficients that exponentiate directly into multiplicative "
        "percentage effects."
    )

    add_para(doc,
        "The model is fitted on N = 2,714 store-days across seven stores and achieves "
        "R² = 0.667 (adjusted R² = 0.664) on the log-transformed target.  In-sample "
        "WAPE on the original scale is 39.5%.  The 28-parameter specification covers "
        "four weather predictors, eleven month dummies, six day-of-week dummies, and "
        "six store fixed effects."
    )

    add_para(doc,
        "Weather effects.  Table 5 reports the four weather coefficient estimates "
        "expressed as percentage changes in daily Ice Cream revenue per unit increase "
        "in each predictor.  Temperature is the dominant weather driver: each "
        "additional degree Celsius is associated with a 5.7% increase in Ice Cream "
        "revenue (95% CI: +4.4% to +7.1%), and each additional hour of daylight with "
        "an 11.2% increase (+4.6% to +18.2%).  Both effects survive BH FDR correction "
        "across the full coefficient family.  Each additional millimetre of precipitation "
        "reduces revenue by 4.7% (-5.4% to -3.9%).  Wind speed shows a small negative "
        "point estimate (-0.4%) but is not significant (p = 0.108) after controlling "
        "for temperature and seasonality."
    )

    # Table 5: Weather % effects
    add_table_caption(doc,
        "Table 5: OLS weather effects on daily Ice Cream revenue.  Coefficients "
        "exponentiated to % change per unit increase.  HC3 robust 95% CIs.  "
        "FDR-adjusted p-values (BH) across all 27 non-intercept coefficients."
    )

    t5_headers = ["Weather predictor", "Unit change", "% Effect",
                  "95% CI", "FDR-adj. p", "Sig."]
    t5_data = [
        ("Temp max",      "+1 C",      "+5.7%",  "[+4.4%, +7.1%]",     "<0.001", "Y"),
        ("Daylight",      "+1 h",      "+11.2%", "[+4.6%, +18.2%]",    "<0.001", "Y"),
        ("Precipitation", "+1 mm",     "-4.7%",  "[-5.4%, -3.9%]",     "<0.001", "Y"),
        ("Wind speed",    "+1 m/s",    "-0.4%",  "[-0.9%, +0.1%]",     "0.108",  ""),
    ]
    t5 = doc.add_table(rows=1 + len(t5_data), cols=6)
    t5.style = "Table Grid"
    C = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(t5_headers):
        _cell(t5.cell(0, j), h, bold=True, size=_FONT_HDR, align=C)
    for i, row in enumerate(t5_data, 1):
        _cell(t5.cell(i, 0), row[0])
        for j in range(1, 6):
            _cell(t5.cell(i, j), row[j], align=C, size=_FONT_BODY)
    _style_header(t5)
    _set_col_widths(t5, [1.30, 0.80, 0.65, 1.30, 0.75, 0.40])
    add_spacer(doc)

    add_para(doc,
        "Calendar and store effects.  Day-of-week effects are large and highly "
        "significant: Saturday carries a +122% premium over Monday (95% CI: "
        "+95% to +152%), Friday +109% (+83% to +139%), and Sunday +61% (+41% to +83%).  "
        "These patterns are consistent with the footfall seasonality described in "
        "Section 2.3.  Store fixed effects show substantial variation relative to "
        "Store A: Stores D, E, and F are 44-50% lower, and Store G 89% lower, "
        "reflecting differences in size, location, and trading format."
    )

    add_figure(doc, "model_coefficient_forest.png",
        "Figure 16: OLS coefficient forest plot.  Left panel: weather % effects "
        "per unit change (95% HC3 CI; grey = not significant at 5%).  Centre: "
        "store fixed effects relative to Store A.  Right: month effects relative "
        "to January.  Error bars are 95% confidence intervals."
    )

    # ── 5.2  XGBoost Forecaster ────────────────────────────────────────────
    doc.add_heading("5.2  XGBoost Forecaster", level=2)

    add_para(doc,
        "An XGBoost gradient-boosting regressor is trained on a richer feature set "
        "that extends the OLS specification with autoregressive lag features: the "
        "previous day's revenue (lag-1), the same weekday from the prior week (lag-7), "
        "and a 7-day rolling mean (all computed within each store in strict date order "
        "to prevent leakage).  Additional calendar features include a weekend flag and "
        "cyclical day-of-year encodings (sine and cosine of the day number, preserving "
        "the circular calendar structure).  The model is validated using "
        "TimeSeriesSplit cross-validation (5 folds) keyed on unique dates so all "
        "stores' observations for a given date are always assigned to the same fold."
    )

    add_para(doc,
        "Table 6 reports per-fold WAPE for the seasonal-naive baseline (lag-7 "
        "forecast) and XGBoost, together with the fold-level R².  Overall, XGBoost "
        "achieves WAPE = 39.5% against a naive baseline of 47.2%, beating the "
        "baseline on four of five folds.  Fold 1 is the exception: with fewer than "
        "three months of training data the model cannot capture full seasonal "
        "variation, and the lag features degrade as the forecast horizon extends "
        "beyond seven days.  XGBoost out-of-sample WAPE is essentially identical "
        "to the OLS in-sample WAPE (39.5% vs 39.5%), which is an honest result: "
        "both methods explain roughly the same variation.  XGBoost's incremental "
        "value lies in its non-linear lag-weather interactions and, as described in "
        "Section 5.3, its capacity to produce calibrated prediction intervals."
    )

    # Table 6: CV metrics
    add_table_caption(doc,
        "Table 6: XGBoost 5-fold time-series cross-validation metrics.  "
        "Naive baseline = lag-7 (seasonal-naive forecast).  Overall row aggregates "
        "all test folds.  Fold 1 note: training window covers fewer than three months."
    )

    t6_headers = ["Fold", "Naive WAPE", "XGB WAPE", "XGB R²", "n test"]
    t6_data = [
        ("1",       "43.9%", "52.8%", "-0.02", "480"),
        ("2",       "50.5%", "32.6%",  "0.59", "483"),
        ("3",       "47.8%", "33.1%",  "0.69", "473"),
        ("4",       "41.8%", "40.3%",  "0.65", "463"),
        ("5",       "48.7%", "37.2%",  "0.47", "405"),
        ("Overall", "47.2%", "39.5%",  "0.55", "2,304"),
    ]
    t6 = doc.add_table(rows=1 + len(t6_data), cols=5)
    t6.style = "Table Grid"
    C = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(t6_headers):
        _cell(t6.cell(0, j), h, bold=True, size=_FONT_HDR, align=C)
    for i, (fold, naive, xgb_w, xgb_r, n) in enumerate(t6_data, 1):
        _cell(t6.cell(i, 0), fold,  align=C, size=_FONT_BODY)
        _cell(t6.cell(i, 1), naive, align=C, size=_FONT_BODY)
        _cell(t6.cell(i, 2), xgb_w, align=C, size=_FONT_BODY)
        _cell(t6.cell(i, 3), xgb_r, align=C, size=_FONT_BODY)
        _cell(t6.cell(i, 4), n,     align=C, size=_FONT_BODY)
    _style_header(t6)
    # Bold the overall row
    for cell in t6.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    _set_col_widths(t6, [0.55, 0.85, 0.85, 0.70, 0.65])
    add_spacer(doc)

    add_figure(doc, "model_cv_metrics.png",
        "Figure 17: Cross-validation WAPE by fold.  Grouped bars show seasonal-naive "
        "(red) vs XGBoost (blue); dashed lines show overall WAPE; dotted green line "
        "shows OLS in-sample WAPE for reference.  XGBoost beats the naive baseline on "
        "4/5 folds; Fold 1 underperforms due to a short early training window."
    )

    add_figure(doc, "model_predicted_vs_actual.png",
        "Figure 18: XGBoost predicted vs actual Ice Cream revenue across all test "
        "folds (revenue indexed, max = 100).  Observations cluster near the diagonal "
        "for mid-range values; variance is higher at the extremes, consistent with "
        "the heteroscedastic residual pattern."
    )

    # ── 5.3  Prediction Intervals ──────────────────────────────────────────
    doc.add_heading("5.3  Prediction Intervals (Conformal Calibration)", level=2)

    add_para(doc,
        "Point forecasts alone are insufficient for operational planning; prediction "
        "intervals are needed to quantify uncertainty.  The XGBoost model produces "
        "raw 90% prediction intervals using separate quantile regressors trained on "
        "the lower and upper quantiles.  Without calibration, the raw intervals "
        "achieve only 64.5% empirical coverage across the test folds -- well below "
        "the 90% target -- because quantile regression in gradient-boosted trees "
        "systematically underestimates uncertainty."
    )

    add_para(doc,
        "Conformalized Quantile Regression (CQR) is applied to correct this.  Within "
        "each training fold, the last 20% of dates form a held-out calibration window.  "
        "Conformity scores are computed on that window as the amount by which the raw "
        "interval must be symmetrically expanded to just cover each observation.  The "
        "finite-sample adjusted 90th percentile of those scores -- the CQR adjustment "
        "q_hat -- is then added to both bounds of the test-set intervals.  CQR is "
        "theoretically guaranteed to achieve at least 90% marginal coverage for "
        "exchangeable data."
    )

    add_para(doc,
        "After CQR calibration, empirical coverage rises from 64.5% to 81.2%, a gain "
        "of 16.7 percentage points.  The remaining 8.8-point shortfall from the 90% "
        "target reflects temporal non-stationarity: the test period can exhibit "
        "different distributional properties from the calibration window, violating "
        "the exchangeability assumption.  The per-fold CQR adjustments varied "
        "substantially, with one fold requiring a notably larger expansion than the "
        "others -- evidence that the data contains regime shifts that a stationary "
        "conformal procedure cannot fully absorb.  Adaptive conformal prediction "
        "(which recalibrates the adjustment at each time step) would be the natural "
        "next step to close the remaining gap."
    )

    add_figure(doc, "model_forecast_intervals.png",
        "Figure 19: Example forecast with CQR-calibrated 90% prediction interval for "
        "Store A, last 60 test days.  Revenue is indexed (max = 100).  The shaded band "
        "shows the calibrated interval; the black line is actual revenue; the dashed "
        "blue line is the XGBoost median forecast.  Local coverage is shown in the "
        "panel title."
    )

    # ── 5.4  Feature Importance (SHAP) ────────────────────────────────────
    doc.add_heading("5.4  Feature Importance (SHAP)", level=2)

    add_para(doc,
        "SHAP (SHapley Additive exPlanations) values are computed on the full dataset "
        "using a tree explainer.  Figure 20 shows mean absolute SHAP values, which "
        "measure each feature's average contribution to the model's output across all "
        "observations."
    )

    add_para(doc,
        "The model is predominantly autoregressive.  The previous day's revenue "
        "(lag-1) is by far the most important single feature, followed by the 7-day "
        "rolling mean.  Together these two lag features dominate the attribution, "
        "which reflects the strong day-to-day persistence and weekly rhythm in ice "
        "cream demand.  Note that lag-1, rolling-7-day mean, and lag-7 are mutually "
        "collinear; their individual SHAP values are ambiguous (collinear features "
        "share attribution depending on which was observed first in the tree splits), "
        "but their combined contribution is correctly measured."
    )

    add_para(doc,
        "Among non-lag features, temperature and daylight are the strongest drivers, "
        "consistent with the OLS regression findings in Section 5.1.  Precipitation "
        "and day-of-week rank next.  The cyclical day-of-year features (sine and "
        "cosine encodings) capture smooth within-year seasonality that the monthly "
        "dummies in the OLS model approximate discretely.  Month, weekend flag, and "
        "wind speed have modest but non-zero importance.  The overall picture is that "
        "weather is a real but secondary driver: the autoregressive lag structure "
        "explains the majority of variance, while weather variables account for a "
        "further material share."
    )

    add_figure(doc, "model_shap_summary.png",
        "Figure 20: SHAP feature importance (mean |SHAP value|, XGBoost trained on "
        "full dataset).  Lag-1 revenue and rolling 7-day mean dominate; temperature "
        "and daylight are the leading weather predictors.  Collinear lag features "
        "share attribution; their individual values should be interpreted jointly."
    )

    # ── 5.5  Caveats ──────────────────────────────────────────────────────
    doc.add_heading("5.5  Caveats", level=2)

    add_para(doc,
        "Four limitations apply to both models.  First, both are trained and validated "
        "on the available historical period; if the business expands to new locations "
        "or product ranges, retraining will be required.  Second, the autoregressive "
        "lag features in the XGBoost model mean that forecast accuracy degrades as "
        "the horizon extends beyond one week (lags become unavailable and must "
        "themselves be forecast).  Third, the conformal calibration assumes approximate "
        "stationarity between the calibration and test windows; the observed shortfall "
        "from 90% coverage indicates some non-stationarity is present.  Fourth, "
        "neither model accounts for causal drivers such as marketing campaigns, price "
        "changes, or public holidays, all of which may confound the weather-revenue "
        "relationship and reduce out-of-sample accuracy if these factors change."
    )


def main() -> None:
    doc = Document(REPORT)
    before_paras   = len(doc.paragraphs)
    before_tables  = len(doc.tables)
    before_images  = len(doc.inline_shapes)

    build_section(doc)
    doc.save(REPORT)

    after_paras   = len(doc.paragraphs)
    after_tables  = len(doc.tables)
    after_images  = len(doc.inline_shapes)

    print(f"Paragraphs : {before_paras} -> {after_paras}  (+{after_paras - before_paras})")
    print(f"Tables     : {before_tables} -> {after_tables}  (+{after_tables - before_tables})")
    print(f"Images     : {before_images} -> {after_images}  (+{after_images - before_images})")

    # Verify headings
    print("\nAll headings:")
    for p in doc.paragraphs:
        if "Heading" in p.style.name and p.text.strip():
            print(f"  {p.style.name}: {p.text}")

    # DKK scan
    import re
    full_text = "\n".join(p.text for p in doc.paragraphs)
    full_text += "\n".join(
        c.text for t in doc.tables
        for row in t.rows for c in row.cells
    )
    hits = re.findall(r"\bDKK\b", full_text, re.IGNORECASE)
    print(f"\nDKK occurrences in new text: {len(hits)}")

    print(f"\nImages intact: {after_images}  (expected >= 15 + 5 = 20)")


if __name__ == "__main__":
    main()
