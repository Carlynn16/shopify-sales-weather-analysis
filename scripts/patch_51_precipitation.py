"""
Patch Section 5.1: insert one clarifying sentence after the precipitation
coefficient sentence, then re-export reports/report.docx to reports/report.pdf.

Run from project root:  python scripts/patch_51_precipitation.py
"""
from pathlib import Path

from docx import Document

DOCX = Path("reports/report.docx")
PDF  = Path("reports/report.pdf")

# The exact substring that ends the precipitation sentence in para 134.
ANCHOR = "-3.9%)."

# The clarifying sentence to insert immediately after ANCHOR.
INSERT = (
    "  This significant conditional effect contrasts with the weaker bivariate "
    "rain–revenue association reported in Section 4.1: by controlling for "
    "temperature and seasonality, the regression isolates precipitation’s "
    "independent contribution to revenue, separating it from the seasonal "
    "co-movement shared with colder, wetter periods."
)


def patch(doc: Document) -> bool:
    """
    Find the precipitation paragraph in Section 5.1 and splice INSERT
    into run[0].text right after ANCHOR.  Returns True if the patch
    was applied, False if ANCHOR was not found.
    """
    in_51 = False
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and "5.1" in p.text:
            in_51 = True
            continue
        if p.style.name.startswith("Heading") and in_51:
            break
        if not in_51:
            continue

        if ANCHOR not in p.text:
            continue

        # Paragraph found.  The text lives in a single run.
        run = p.runs[0]
        idx = run.text.index(ANCHOR) + len(ANCHOR)
        run.text = run.text[:idx] + INSERT + run.text[idx:]
        return True

    return False


def export_pdf() -> int:
    from docx2pdf import convert
    convert(str(DOCX), str(PDF))
    return PDF.stat().st_size // 1024


def main() -> None:
    doc = Document(DOCX)

    applied = patch(doc)
    if not applied:
        raise RuntimeError(f"Anchor {ANCHOR!r} not found — nothing changed.")

    doc.save(DOCX)
    print("Patch applied.  Verifying inserted text...")

    # Reload and confirm
    doc2 = Document(DOCX)
    for p in doc2.paragraphs:
        if "independent contribution to revenue" in p.text:
            # Print just the precipitation + new sentence portion
            start = p.text.find("Each additional millimetre")
            snippet = p.text[start:start + 320] if start != -1 else p.text[:320]
            print(f"  ✔ Found:\n    {snippet}\n")
            break
    else:
        raise RuntimeError("Inserted sentence not found after save — check document.")

    print("Re-exporting PDF...")
    kb = export_pdf()
    print(f"  ✔ PDF written: {PDF}  ({kb:,} KB)")


if __name__ == "__main__":
    main()
