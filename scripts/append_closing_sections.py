"""
Append closing sections to reports/report.docx:
  - Executive Summary  (inserted as front matter, before Section 1)
  - 6  Recommendations
  - 7  Limitations & Next Steps

Run from project root:  python scripts/append_closing_sections.py
All figures are scale-free; no DKK, no absolute q-hat values.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

REPORT    = Path("reports/report.docx")
_FONT_BODY = Pt(9)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _add_run(para, text: str, bold: bool = False) -> None:
    r = para.add_run(text)
    r.bold = bold
    r.font.size = _FONT_BODY


def add_para(doc, text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    _add_run(p, text)
    p.paragraph_format.space_after = Pt(6)


def add_bold_lead(doc, lead: str, rest: str, style: str = "List Bullet") -> None:
    """Bullet paragraph with a bold lead phrase followed by normal text."""
    p = doc.add_paragraph(style=style)
    _add_run(p, lead, bold=True)
    _add_run(p, rest)


def add_spacer(doc) -> None:
    doc.add_paragraph(style="Normal").paragraph_format.space_after = Pt(6)


def add_page_break(doc) -> None:
    p = doc.add_paragraph(style="Normal")
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Section content
# ---------------------------------------------------------------------------

def build_exec_summary(doc) -> None:
    add_page_break(doc)
    doc.add_heading("Executive Summary", level=1)

    add_para(doc,
        "This report analyses 14 months of point-of-sale transaction data from a "
        "Danish artisan food chain operating seven physical stores, all running on "
        "Shopify.  The dataset covers approximately 454,000 transaction lines across "
        "around 280,000 completed orders.  The analysis pursues two business goals: "
        "understanding what drives revenue across products, categories, stores, and "
        "seasons; and building a weather-driven forecast to support staffing, "
        "ice-cream preparation, and inventory decisions."
    )

    add_para(doc,
        "Revenue is highly concentrated: roughly 20 products account for 80% of "
        "total revenue, with Ice Cream the dominant category at approximately 57%.  "
        "Sales follow a strong seasonal pattern, peaking in May and June, with "
        "Friday and Saturday the highest-revenue days and the 14:00–16:00 window "
        "the busiest part of the trading day.  Weather effects on daily Ice Cream "
        "revenue are large and statistically significant: each additional degree "
        "Celsius is associated with approximately 5.7% higher revenue (95% CI: "
        "+4.4% to +7.1%) and each additional hour of daylight with approximately "
        "11.2% more (+4.6% to +18.2%), after controlling for day-of-week, month, "
        "and store fixed effects.  A data-quality investigation identified that one "
        "store’s POS terminal was not synced to Shopify product codes, causing "
        "a concentrated pattern of unattributed transactions; the issue declined "
        "sharply following a POS reconfiguration."
    )

    add_para(doc,
        "An XGBoost forecaster, validated with five-fold time-series "
        "cross-validation, achieves a weighted absolute percentage error (WAPE) of "
        "39.5% against a seasonal-naive baseline of 47.2% — a 7.7-percentage-point "
        "improvement.  The model is primarily autoregressive, with weather variables "
        "providing a further material contribution.  Calibrated 90% prediction "
        "intervals reach approximately 81% empirical coverage after conformal "
        "calibration, with the remaining gap attributable to distributional shifts "
        "across time."
    )

    add_para(doc,
        "The three most actionable recommendations are: (1) concentrate staffing "
        "and ice-cream preparation on Friday, Saturday, and the 14:00–16:00 "
        "window, where revenue is disproportionately high; (2) use the "
        "weather-driven forecast to scale ice-cream prep and staffing ahead of "
        "warm, long-daylight days — a forecast of warmer, longer days is a direct "
        "signal to increase preparation; and (3) fix the POS configuration at the "
        "affected store to restore full per-category revenue attribution."
    )


def build_recommendations(doc) -> None:
    doc.add_heading("6  Recommendations", level=1)

    add_para(doc,
        "Five concrete recommendations follow directly from the analysis.  Each is "
        "grounded in a specific quantitative finding."
    )

    add_bold_lead(doc,
        "Concentrate staffing and preparation on peak periods.  ",
        "Friday and Saturday carry revenue premiums of approximately +109% and "
        "+122% over Monday respectively.  Within the trading day the 14:00–16:00 "
        "window is the revenue peak.  Scheduling staff and ice-cream preparation to "
        "align with these peaks will reduce both under-service on busy days and "
        "over-staffing on quiet ones."
    )

    add_bold_lead(doc,
        "Use the weather-driven forecast to scale ice-cream demand.  ",
        "Each additional degree Celsius above the baseline is associated with "
        "approximately 5.7% higher daily Ice Cream revenue; each additional hour of "
        "daylight with approximately 11.2%.  The XGBoost forecaster operationalises "
        "these effects: a warm, long-daylight forecast is a direct signal to "
        "increase ice-cream stock, preparation capacity, and staffing.  Summer and "
        "warm/long-daylight days drive the largest upswings and carry the highest "
        "risk of under-preparation."
    )

    add_bold_lead(doc,
        "Build a counter-seasonal product mix.  ",
        "Ice Cream revenue drops sharply in winter.  Hot Beverages and "
        "Chocolate/Christmas products show the inverse seasonal profile and can "
        "partially offset the off-season trough.  Cross-training staff across "
        "seasonal product lines will reduce scheduling friction as the product mix "
        "shifts between summer and winter."
    )

    add_bold_lead(doc,
        "Fix the POS configuration at the affected store.  ",
        "One store’s POS terminal was not synced to Shopify product codes, "
        "generating a concentrated block of transactions with no product "
        "attribution.  This understates that store’s per-category revenue and "
        "distorts any category-level comparisons involving it.  Completing the POS "
        "reconfiguration (already partially enacted) and back-filling historical "
        "data where possible will restore accurate attribution."
    )

    add_bold_lead(doc,
        "Focus inventory and promotion on the core product set.  ",
        "Approximately 20 products drive 80% of total revenue — a clean Pareto "
        "distribution.  Inventory depth, promotional spend, and operational "
        "attention should be weighted towards these lines; long-tail products with "
        "negligible revenue contribution are candidates for rationalisation."
    )

    add_spacer(doc)


def build_limitations(doc) -> None:
    doc.add_heading("7  Limitations & Next Steps", level=1)

    add_para(doc,
        "Five limitations apply to the analysis and models presented in this report."
    )

    add_bold_lead(doc,
        "Observational data — no causal claims.  ",
        "The dataset is an observational transaction record; no controlled "
        "experiment was run.  All stated effects (weather on revenue, day-of-week "
        "premiums) are associations, not causal estimates.  Unmeasured confounders "
        "— marketing campaigns, pricing changes, local events — may drive part "
        "of the observed variation.  Causal conclusions would require experimental "
        "or quasi-experimental designs.",
        style="List Bullet"
    )

    add_bold_lead(doc,
        "Temporal autocorrelation.  ",
        "Daily revenue observations from the same store are serially correlated.  "
        "The bootstrap confidence intervals reported for group comparisons and "
        "store-level temperature sensitivities are more reliable than the "
        "conventional asymptotic standard errors, because they do not assume "
        "independence across observations.  Raw p-values from parametric tests "
        "should be treated as indicative rather than definitive.",
        style="List Bullet"
    )

    add_bold_lead(doc,
        "Severe collinearity among weather predictors.  ",
        "Temperature, feels-like temperature, daylight duration, and calendar "
        "season are highly intercorrelated.  The regression uses VIF-pruned "
        "predictors (feels-like excluded) to reduce this, but the retained "
        "coefficients still carry shared attribution with the excluded variables.  "
        "Individual predictor effects should be interpreted jointly, not in "
        "isolation.",
        style="List Bullet"
    )

    add_bold_lead(doc,
        "The forecaster is predominantly autoregressive.  ",
        "The XGBoost model’s SHAP analysis shows that recent lag revenue "
        "(previous day, 7-day rolling mean) dominates feature importance; weather "
        "variables are a secondary but material contributor.  Forecast accuracy "
        "therefore degrades as the horizon extends beyond one week, because lag "
        "features must themselves be forecast.  The model is best suited to "
        "short-horizon (1–7 day) operational planning rather than long-range "
        "demand projection.",
        style="List Bullet"
    )

    add_bold_lead(doc,
        "Conformal prediction interval coverage is capped near 81%.  ",
        "The CQR-calibrated 90% prediction intervals achieve approximately 81% "
        "empirical coverage — 9 percentage points below the nominal target.  "
        "The shortfall reflects temporal non-stationarity between the calibration "
        "and test windows: the exchangeability assumption required for exact "
        "conformal coverage guarantees is violated when the data distribution "
        "shifts over time.  The q-hat adjustment varies substantially across folds, "
        "confirming the presence of regime shifts that a stationary conformal "
        "procedure cannot fully absorb.",
        style="List Bullet"
    )

    add_spacer(doc)

    doc.add_heading("Next Steps", level=2)

    add_bold_lead(doc,
        "Adaptive conformal prediction.  ",
        "Replacing the static CQR calibration with an adaptive conformal procedure "
        "— which recalibrates the interval adjustment at each time step using a "
        "rolling window — would close the remaining coverage gap by tracking "
        "distributional shifts in real time.",
        style="List Bullet"
    )

    add_bold_lead(doc,
        "Incorporate external events.  ",
        "Neither model accounts for marketing campaigns, price changes, or public "
        "holidays.  Adding a public-holiday indicator and, where available, "
        "campaign flags would reduce residual variance and improve both the "
        "regression coefficient estimates and the forecaster’s out-of-sample "
        "accuracy.",
        style="List Bullet"
    )

    add_bold_lead(doc,
        "Expand the store coverage.  ",
        "Both models were trained and validated on the current seven-store network.  "
        "If the business expands to new locations, the models should be retrained "
        "on data that includes the new stores; extrapolating store fixed effects to "
        "novel locations is unreliable.",
        style="List Bullet"
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    doc = Document(REPORT)

    n_paras_before  = len(doc.paragraphs)
    n_tables_before = len(doc.tables)
    n_images_before = len(doc.inline_shapes)

    # ── 1. Locate the Section 1 heading element BEFORE adding anything ──────
    # Para 15 is "1  Data & Methodology" in the current document.
    # We capture the XML element reference; it stays valid through all edits.
    section1_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and p.text.strip().startswith("1 "):
            section1_idx = i
            break
    if section1_idx is None:
        sys.exit("ERROR: Could not locate '1  Data & Methodology' heading.")

    section1_elem = doc.paragraphs[section1_idx]._element

    # ── 2. Append Recommendations and Limitations (normal append) ────────────
    build_recommendations(doc)
    build_limitations(doc)

    # ── 3. Build Executive Summary (temporarily appended) ───────────────────
    # Record how many paragraphs exist after step 2.
    n_after_closing = len(doc.paragraphs)

    build_exec_summary(doc)

    # All paragraphs added in step 3 are currently at the end of body.
    # Collect their XML elements.
    exec_elems = [p._element for p in doc.paragraphs[n_after_closing:]]

    # ── 4. Move Executive Summary to just before Section 1 ──────────────────
    # lxml's addprevious() moves the element from its current position to
    # immediately before section1_elem each time.  Iterating forward keeps
    # the final order correct: each element leapfrogs the previous one so the
    # sequence page_break → heading → para1 → … → para4 → section1 is built
    # up correctly from left to right.
    for elem in exec_elems:
        section1_elem.addprevious(elem)

    # ── 5. Save ──────────────────────────────────────────────────────────────
    doc.save(REPORT)

    # ── 6. Verification ──────────────────────────────────────────────────────
    doc2 = Document(REPORT)
    n_paras_after  = len(doc2.paragraphs)
    n_tables_after = len(doc2.tables)
    n_images_after = len(doc2.inline_shapes)

    print(f"Paragraphs : {n_paras_before} -> {n_paras_after}  (+{n_paras_after - n_paras_before})")
    print(f"Tables     : {n_tables_before} -> {n_tables_after}  (+{n_tables_after - n_tables_before})")
    print(f"Images     : {n_images_before} -> {n_images_after}  (+{n_images_after - n_images_before})")

    print("\nAll headings in order:")
    for p in doc2.paragraphs:
        if "Heading" in p.style.name and p.text.strip():
            print(f"  {p.style.name}: {p.text}")

    import re
    full_text = "\n".join(p.text for p in doc2.paragraphs)
    full_text += "\n".join(
        c.text for t in doc2.tables
        for row in t.rows for c in row.cells
    )

    suspects = {
        "DKK": r"\bDKK\b",
        "Vedbaek/Vedb": r"Vedb",
        "Hellerup": r"Hellerup",
        "Gentofte": r"Gentofte",
        "Lyngby": r"Lyngby",
    }
    print("\nSafety scan:")
    for label, pat in suspects.items():
        hits = re.findall(pat, full_text, re.IGNORECASE)
        status = "FAIL" if hits else "ok"
        print(f"  [{status}] {label}: {len(hits)}")

    print(f"\nImages intact: {n_images_after}  (expected 20)")

    # Print the Executive Summary text for inspection
    print("\n--- Executive Summary text ---")
    in_exec = False
    for p in doc2.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == "Executive Summary":
            in_exec = True
        elif p.style.name == "Heading 1" and in_exec:
            break
        if in_exec and p.text.strip():
            print(p.text)
            print()


if __name__ == "__main__":
    import sys
    main()
