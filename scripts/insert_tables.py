"""
Insert four results tables into Section 4 of report.docx.

Insertion points are located by anchor text so the script is robust to
surrounding edits.  All existing paragraphs, figures, and manual edits
are untouched.
"""
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

REPORT = Path("report.docx")

# Column widths (inches) per table
_W_CORR   = [0.85, 1.05, 0.60, 0.60, 1.25, 0.75, 0.35]   # 7 cols = 5.45 in
_W_VIF    = [2.20, 1.40]                                     # 2 cols = 3.60 in
_W_GROUP  = [1.40, 0.90, 0.90, 1.35, 0.75]                  # 5 cols = 5.30 in
_W_SENS   = [0.75, 0.75, 1.35, 0.75]                        # 4 cols = 3.60 in

_HEADER_BG  = "D9E1F2"   # pale blue
_FONT_HDR   = Pt(9)
_FONT_BODY  = Pt(9)
_FONT_CAP   = Pt(10)

_VAR_LABELS = {
    "temperature_2m_max":       "Temp max",
    "apparent_temperature_max": "App. temp max",
    "daylight_duration":        "Daylight",
    "precipitation_sum":        "Precipitation",
    "rain_sum":                 "Rain",
    "snowfall_sum":             "Snowfall",
    "windspeed_10m_max":        "Wind speed",
}


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _fp(v):
    """Format p-value."""
    if pd.isna(v):
        return "n/a"
    if v < 0.001:
        return "<0.001"
    return f"{v:.3f}"


def _fr(v, dp=2):
    if pd.isna(v):
        return "--"
    return f"{v:+.{dp}f}"


def _fci(lo, hi, dp=2):
    if pd.isna(lo) or pd.isna(hi):
        return "--"
    return f"[{lo:+.{dp}f}, {hi:+.{dp}f}]"


def _fvif(v):
    if pd.isna(v):
        return "--"
    if v >= 10000:
        return f"{v:,.0f}"
    if v >= 100:
        return f"{v:,.1f}"
    return f"{v:.2f}"


def _sig(p_adj):
    return "Y" if (pd.notna(p_adj) and p_adj < 0.05) else ""


# ---------------------------------------------------------------------------
# Table styling helpers
# ---------------------------------------------------------------------------

def _set_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    tcPr.append(shd)


def _cell(cell, text, bold=False, size=_FONT_BODY,
          align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run = para.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    run.font.size = size
    return run


def _set_col_widths(table, widths_in):
    for row in table.rows:
        for i, w in enumerate(widths_in):
            row.cells[i].width = Inches(w)


def _style_header(table):
    for cell in table.rows[0].cells:
        _set_bg(cell, _HEADER_BG)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold      = True
                run.font.size = _FONT_HDR


# ---------------------------------------------------------------------------
# Insertion utility
# ---------------------------------------------------------------------------

def _insert_after(doc, anchor_para, caption_text, table):
    """
    Place  caption_text | table  immediately after anchor_para.
    Final order: ... anchor | caption | table | spacer | ...
    """
    # 1. Caption paragraph (currently at doc end)
    cap = doc.add_paragraph(style="Normal")
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after  = Pt(2)
    run = cap.add_run(caption_text)
    run.bold      = True
    run.font.size = _FONT_CAP

    # 2. Move table to right after anchor, then caption between them
    anchor_para._p.addnext(table._tbl)
    anchor_para._p.addnext(cap._p)

    # 3. Spacer after table
    spacer = doc.add_paragraph(style="Normal")
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(6)
    table._tbl.addnext(spacer._p)


def _find_para(doc, snippet, start=0):
    """First paragraph at index >= start whose .text contains snippet."""
    for i, p in enumerate(doc.paragraphs):
        if i < start and snippet not in p.text:
            continue
        if snippet in p.text:
            return i, p
    return None, None


# ---------------------------------------------------------------------------
# Table builders
# ---------------------------------------------------------------------------

def _build_corr_table(doc, corr_df):
    """
    Table 1: Weather-Revenue Correlations
    Cols: Outcome | Predictor | Pearson r | Spearman r | 95% CI | FDR-adj. p | Sig.
    """
    outcomes = ["Total", "Ice Cream", "Hot Beverages"]
    headers  = ["Outcome", "Predictor",
                "Pearson r", "Spearman r", "95% CI (Spearman)", "FDR-adj. p", "Sig."]

    rows = []
    for outcome in outcomes:
        sub = corr_df[corr_df["outcome"] == outcome]
        first = True
        for _, r in sub.iterrows():
            rows.append({
                "outcome":   outcome if first else "",
                "predictor": _VAR_LABELS.get(r["weather_var"], r["weather_var"]),
                "pear_r":    _fr(r["pearson_r"]),
                "spear_r":   _fr(r["spearman_r"]),
                "ci":        _fci(r["spearman_ci_lower"], r["spearman_ci_upper"]),
                "p_adj":     _fp(r["spearman_p_adj"]),
                "sig":       _sig(r["spearman_p_adj"]),
            })
            first = False

    t = doc.add_table(rows=1 + len(rows), cols=7)
    t.style = "Table Grid"

    C = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(headers):
        _cell(t.cell(0, j), h, bold=True, size=_FONT_HDR, align=C)

    for i, row in enumerate(rows, 1):
        _cell(t.cell(i, 0), row["outcome"],   size=_FONT_BODY)
        _cell(t.cell(i, 1), row["predictor"], size=_FONT_BODY)
        _cell(t.cell(i, 2), row["pear_r"],    size=_FONT_BODY, align=C)
        _cell(t.cell(i, 3), row["spear_r"],   size=_FONT_BODY, align=C)
        _cell(t.cell(i, 4), row["ci"],        size=_FONT_BODY, align=C)
        _cell(t.cell(i, 5), row["p_adj"],     size=_FONT_BODY, align=C)
        _cell(t.cell(i, 6), row["sig"],       size=_FONT_BODY, align=C)

    _style_header(t)
    _set_col_widths(t, _W_CORR)
    return t


def _build_vif_table(doc, vif_df):
    """Table 2: VIF — Predictor | VIF"""
    t = doc.add_table(rows=1 + len(vif_df), cols=2)
    t.style = "Table Grid"

    C = WD_ALIGN_PARAGRAPH.CENTER
    _cell(t.cell(0, 0), "Weather Predictor", bold=True, size=_FONT_HDR)
    _cell(t.cell(0, 1), "VIF",               bold=True, size=_FONT_HDR, align=C)

    for i, row in vif_df.iterrows():
        _cell(t.cell(i + 1, 0), _VAR_LABELS.get(row["predictor"], row["predictor"]))
        _cell(t.cell(i + 1, 1), _fvif(row["VIF"]),
              align=WD_ALIGN_PARAGRAPH.RIGHT, size=_FONT_BODY)

    _style_header(t)
    _set_col_widths(t, _W_VIF)
    return t


def _build_group_table(doc, group_df):
    """Table 3: Group Comparisons"""
    # Sort by rank-biserial r descending
    gdf = group_df.sort_values("rank_biserial_r", ascending=False).reset_index(drop=True)

    headers = ["Split", "Rank-biserial r", "Median diff (pp)",
               "95% CI", "FDR-adj. p"]
    t = doc.add_table(rows=1 + len(gdf), cols=5)
    t.style = "Table Grid"

    C = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(headers):
        _cell(t.cell(0, j), h, bold=True, size=_FONT_HDR, align=C)

    for i, row in gdf.iterrows():
        _cell(t.cell(i + 1, 0), row["split"])
        _cell(t.cell(i + 1, 1), _fr(row["rank_biserial_r"]), align=C, size=_FONT_BODY)
        _cell(t.cell(i + 1, 2), _fr(row["median_diff"], dp=3), align=C, size=_FONT_BODY)
        _cell(t.cell(i + 1, 3), _fci(row["ci_lower"], row["ci_upper"], dp=3),
              align=C, size=_FONT_BODY)
        _cell(t.cell(i + 1, 4), _fp(row["p_adj"]), align=C, size=_FONT_BODY)

    _style_header(t)
    _set_col_widths(t, _W_GROUP)
    return t


def _build_sens_table(doc, sens_df):
    """Table 4: Per-Store Temperature Sensitivity"""
    headers = ["Store", "Spearman r", "95% CI", "FDR-adj. p"]
    t = doc.add_table(rows=1 + len(sens_df), cols=4)
    t.style = "Table Grid"

    C = WD_ALIGN_PARAGRAPH.CENTER
    for j, h in enumerate(headers):
        _cell(t.cell(0, j), h, bold=True, size=_FONT_HDR, align=C)

    for i, row in sens_df.iterrows():
        _cell(t.cell(i + 1, 0), row["store_label"])
        _cell(t.cell(i + 1, 1), _fr(row["spearman_r"]), align=C, size=_FONT_BODY)
        _cell(t.cell(i + 1, 2), _fci(row.get("ci_lower"), row.get("ci_upper")),
              align=C, size=_FONT_BODY)
        _cell(t.cell(i + 1, 3), _fp(row["p_adj"]), align=C, size=_FONT_BODY)

    _style_header(t)
    _set_col_widths(t, _W_SENS)
    return t


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    # ── Load stats ─────────────────────────────────────────────────────────
    from src.categorization import categorize
    from src.cleaning import clean_transactions
    from src.config import DATA_DIR
    from src.data_loading import build_transactions, load_raw_tables
    from src.weather import build_weather_panel
    from src.weather_stats import (
        apply_fdr,
        build_daily_category_revenue,
        group_comparisons,
        store_weather_sensitivity,
        weather_correlations,
        weather_multicollinearity,
    )

    print("Loading data and running stats (uses weather cache)...")
    tables     = load_raw_tables(DATA_DIR)
    tx         = categorize(clean_transactions(build_transactions(tables)))
    locs       = tables["locations"]
    store_panel, daily_panel = build_weather_panel(tx, locs)

    daily_cat  = build_daily_category_revenue(tx)
    corr_df    = weather_correlations(daily_panel, daily_cat)
    _, vif_df  = weather_multicollinearity(daily_panel)
    group_df   = group_comparisons(daily_panel)
    sens_df    = store_weather_sensitivity(store_panel)

    # ── Open document ───────────────────────────────────────────────────────
    doc = Document(REPORT)
    before_paras  = len(doc.paragraphs)
    before_images = len(doc.inline_shapes)
    before_tables = len(doc.tables)

    # ── Locate anchors ──────────────────────────────────────────────────────
    # T1: after the "Block-bootstrap..." para (para 89 in original), before Figure 13
    _, anchor1 = _find_para(doc, "Block-bootstrap confidence intervals use 7-day")
    # T2: after "must be collapsed to a single measure" para (para 96), before H2 4.3
    _, anchor2 = _find_para(doc, "must be collapsed to a single measure")
    # T3: after the Rainy bullet (para 103), before Figure 15
    _, anchor3 = _find_para(doc, "Rainy (> 1 mm) vs Dry: rank-biserial")
    # T4: after the store sensitivity prose (para 107), before H2 4.5
    _, anchor4 = _find_para(doc, "consistent with the category-level findings in Section 4.1")

    assert anchor1, "Anchor 1 not found"
    assert anchor2, "Anchor 2 not found"
    assert anchor3, "Anchor 3 not found"
    assert anchor4, "Anchor 4 not found"

    # ── Build tables ────────────────────────────────────────────────────────
    t1 = _build_corr_table(doc, corr_df)
    t2 = _build_vif_table(doc, vif_df)
    t3 = _build_group_table(doc, group_df)
    t4 = _build_sens_table(doc, sens_df)

    # ── Insert  (reverse order so earlier anchors stay valid) ───────────────
    # Insert T4 first (deepest in doc), then T3, T2, T1
    _insert_after(doc, anchor4,
                  "Table 4: Per-Store Temperature Sensitivity -- "
                  "Spearman correlation with daily maximum temperature "
                  "(sorted most to least sensitive; all FDR-adjusted p < 0.001).",
                  t4)
    _insert_after(doc, anchor3,
                  "Table 3: Group Comparisons -- Mann-Whitney U test results. "
                  "Sorted by effect size (rank-biserial r). "
                  "Median diff is in daily revenue-share percentage points (pp). "
                  "Rainy vs Dry 95% CI crosses zero.",
                  t3)
    _insert_after(doc, anchor2,
                  "Table 2: Variance Inflation Factors for the seven weather "
                  "predictors. Temperature and apparent temperature (VIF ~94) "
                  "and the precipitation cluster (VIF >57,000) are highly "
                  "collinear and cannot enter a regression together.",
                  t2)
    _insert_after(doc, anchor1,
                  "Table 1: Weather-Revenue Correlations -- Pearson r, Spearman r, "
                  "block-bootstrap 95% CI (7-day blocks), and BH FDR-adjusted p-value "
                  "for each outcome x predictor pair. "
                  "Sig. = Y where adjusted p < 0.05.",
                  t1)

    # ── Save ────────────────────────────────────────────────────────────────
    doc.save(REPORT)

    after_paras  = len(doc.paragraphs)
    after_images = len(doc.inline_shapes)
    after_tables = len(doc.tables)

    print(f"\nDocument saved.")
    print(f"  Paragraphs : {before_paras} -> {after_paras}  "
          f"(+{after_paras - before_paras})")
    print(f"  Images     : {before_images} -> {after_images}  "
          f"(+{after_images - before_images} -- should be 0)")
    print(f"  Tables     : {before_tables} -> {after_tables}  "
          f"(+{after_tables - before_tables} -- should be 4)")

    # ── Verify: headings and images intact ─────────────────────────────────
    print("\nAll headings:")
    for p in doc.paragraphs:
        if "Heading" in p.style.name and p.text.strip():
            print(f"  {p.style.name}: {p.text}")

    print(f"\nInline shapes (images): {after_images}  "
          f"{'OK -- 15' if after_images == 15 else 'WARNING: expected 15'}")

    # ── Text dump of the four tables ────────────────────────────────────────
    print("\n" + "=" * 72)
    print("TABLE DUMPS")
    print("=" * 72)
    for t_idx, label in enumerate(
        ["Table 1: Weather-Revenue Correlations",
         "Table 2: VIF",
         "Table 3: Group Comparisons",
         "Table 4: Store Sensitivity"],
        1
    ):
        tbl = doc.tables[t_idx - 1]
        print(f"\n--- {label} ---")
        for row in tbl.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            print("  |  ".join(cells))


if __name__ == "__main__":
    main()
