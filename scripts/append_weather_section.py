"""
Append 'Section 4: Weather Analysis' to reports/report.docx.
Run from the project root: python scripts/append_weather_section.py
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

REPORT = Path("reports/report.docx")
FIGURES = Path("figures")
W = Inches(6.3)   # full content width (8.5 - 2 x 1.1 margins)


def add_para(doc, text, style="Normal"):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_figure(doc, filename, caption):
    """Embed a figure centred, then add a caption paragraph."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURES / filename), width=W)
    cap = doc.add_paragraph(style="Normal")
    cap.add_run(caption)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)


def main():
    doc = Document(REPORT)
    paras_before = len(doc.paragraphs)
    images_before = len(doc.inline_shapes)

    # ── Section 4 ──────────────────────────────────────────────────────────────
    doc.add_heading("4  Weather Analysis", level=1)

    add_para(doc,
        "This section examines the statistical association between daily weather "
        "conditions and daily revenue at the aggregate and category level.  Weather "
        "data were fetched from the Open-Meteo archive API at daily resolution for "
        "each store city and area-averaged to produce a single daily weather panel.  "
        "Seven predictors were examined: maximum temperature, apparent maximum "
        "temperature, daylight duration, total precipitation, rain, snowfall, and "
        "maximum wind speed.  All revenue figures are expressed as daily percentage "
        "shares of total revenue."
    )

    # 4.1 -----------------------------------------------------------------------
    doc.add_heading("4.1  Weather-Revenue Correlations", level=2)

    add_para(doc,
        "Daily total revenue is strongly positively associated with temperature and "
        "daylight.  Spearman correlations with total daily revenue share are r = 0.70 "
        "for maximum temperature (block-bootstrap 95% CI [0.60, 0.77]) and r = 0.70 "
        "for daylight duration (CI [0.58, 0.75]).  The association is considerably "
        "stronger for Ice Cream, the dominant revenue category: Spearman r = 0.82 "
        "with daylight duration (CI [0.75, 0.84]) and r = 0.77 with maximum temperature "
        "(CI [0.67, 0.82]).  Hot Beverages run in the opposite direction -- r = -0.34 "
        "with daylight (CI [-0.52, -0.13]) and r = -0.22 with maximum temperature "
        "(CI [-0.41, -0.03]) -- consistent with customers substituting warm drinks for "
        "cold ones in cooler months."
    )

    add_para(doc,
        "All temperature and daylight associations survive Benjamini-Hochberg false "
        "discovery rate (FDR) correction across the full family of 42 tests (adjusted "
        "p < 0.001).  Snowfall shows a modest negative association with total revenue "
        "(Spearman r = -0.29, CI [-0.37, -0.16], adjusted p < 0.001).  Rain's "
        "association is weak (Spearman r = -0.07) and does not survive FDR correction "
        "(adjusted p = 0.19); the direction is therefore not reliable.  Pearson and "
        "Spearman coefficients agree closely throughout, suggesting the relationships "
        "are not artefacts of non-linearity."
    )

    add_para(doc,
        "Block-bootstrap confidence intervals use 7-day consecutive blocks to preserve "
        "the autocorrelation structure of the daily revenue series; standard p-values "
        "are approximate under temporal autocorrelation and should be read alongside "
        "the bootstrap intervals."
    )

    add_figure(
        doc,
        "weather_revenue_correlations.png",
        "Figure 13: Spearman correlation heatmap -- daily revenue share (Total, "
        "Ice Cream, Hot Beverages) vs seven weather predictors.  Cell values show "
        "Spearman r to two decimal places; asterisks (*) mark associations surviving "
        "BH FDR correction (adjusted p < 0.05).  Temperature and daylight are strongly "
        "positive for Total and Ice Cream revenue, and negative for Hot Beverages.",
    )

    add_figure(
        doc,
        "category_weather_scatter.png",
        "Figure 14: Daily revenue share vs maximum temperature for Ice Cream (left, "
        "blue) and Hot Beverages (right, red), with OLS fit lines.  The opposing "
        "slopes illustrate seasonal substitution between the two categories.  Pearson "
        "r and p-value are shown in each panel legend.",
    )

    # 4.2 -----------------------------------------------------------------------
    doc.add_heading("4.2  Multicollinearity Among Weather Predictors", level=2)

    add_para(doc,
        "The Pearson correlation between maximum temperature and apparent maximum "
        "temperature is r = 0.988; their variance inflation factors (VIF) are 93.6 "
        "and 94.1 respectively.  In practice only one of these variables can be "
        "included in a regression specification; maximum temperature is the preferred "
        "choice as the more physically interpretable measure.  Daylight duration is "
        "also strongly correlated with temperature (r = 0.80) but has a comparatively "
        "modest VIF of 2.78, so it can enter a model alongside temperature provided "
        "apparent temperature is excluded."
    )

    add_para(doc,
        "The precipitation cluster presents a more severe problem: total precipitation "
        "equals rain plus snowfall by construction, producing VIFs of 1,341,169 "
        "(precipitation), 1,238,813 (rain), and 57,888 (snowfall).  These three "
        "variables cannot be interpreted independently in any multivariate model and "
        "must be collapsed to a single measure or replaced with mutually exclusive "
        "flags.  Wind speed (VIF 2.60) is independent of the remaining predictors and "
        "can be included without inflation.  These collinearity findings motivate "
        "pruning the predictor set before fitting the regression model."
    )

    # 4.3 -----------------------------------------------------------------------
    doc.add_heading("4.3  Group Comparisons", level=2)

    add_para(doc,
        "Four binary splits were tested using Mann-Whitney U tests with rank-biserial "
        "correlation as the effect-size measure and 1,000-iteration bootstrap 95% "
        "confidence intervals on the difference in group medians, expressed in daily "
        "revenue-share percentage points (pp).  All four adjusted p-values passed the "
        "BH FDR threshold."
    )

    add_para(doc,
        "Revenue is substantially higher on warm, long-daylight days and in summer.  "
        "Three of the four splits produce large, highly significant effects:"
    )

    for bullet in [
        (
            "Summer (Apr-Sep) vs Winter (Oct-Mar): rank-biserial r = 0.86, median "
            "difference +0.21 pp, 95% CI [0.187, 0.227], adjusted p < 0.001."
        ),
        (
            "Long daylight (>= 11.2 h) vs Short daylight (< 11.2 h): rank-biserial "
            "r = 0.82, median difference +0.19 pp, 95% CI [0.168, 0.218], "
            "adjusted p < 0.001."
        ),
        (
            "Warm (>= 10.8 C) vs Cold (< 10.8 C): rank-biserial r = 0.71, median "
            "difference +0.17 pp, 95% CI [0.156, 0.205], adjusted p < 0.001."
        ),
        (
            "Rainy (> 1 mm) vs Dry: rank-biserial r = -0.12, median difference "
            "-0.026 pp, 95% CI [-0.065, +0.017], adjusted p = 0.034.  The bootstrap "
            "CI crosses zero, so the direction of the precipitation effect is "
            "uncertain despite the test reaching nominal significance."
        ),
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    add_figure(
        doc,
        "group_comparison_boxplots.png",
        "Figure 15: Violin plots of daily revenue share by group for the four binary "
        "splits.  The inner box-and-whisker marks the median and interquartile range.  "
        "The warm/cold, long/short-daylight, and summer/winter splits all show large "
        "revenue shifts; the rainy/dry split shows minimal separation with overlapping "
        "distributions.",
    )

    # 4.4 -----------------------------------------------------------------------
    doc.add_heading("4.4  Per-Store Temperature Sensitivity", level=2)

    add_para(doc,
        "All seven stores show positive Spearman correlations between daily revenue "
        "share and maximum temperature, with all FDR-adjusted p-values below 0.001.  "
        "Store A is the most weather-sensitive (r = 0.75), followed by Store F "
        "(r = 0.69), Store C (r = 0.67), Store E (r = 0.63), Store B (r = 0.59), "
        "Store D (r = 0.43), and Store G (r = 0.31).  The variation likely reflects "
        "differences in store size, footfall mix, and category composition: stores "
        "with a higher ice cream revenue share are expected to show stronger positive "
        "temperature correlations, consistent with the category-level findings in "
        "Section 4.1."
    )

    # 4.5 -----------------------------------------------------------------------
    doc.add_heading("4.5  Caveats", level=2)

    add_para(doc,
        "Three limitations should be noted.  First, all findings are observational: "
        "temporal autocorrelation inflates the effective sample size implied by "
        "standard p-values, so the block-bootstrap confidence intervals (7-day blocks) "
        "are the more reliable uncertainty indicators.  Second, the extreme collinearity "
        "among temperature and apparent temperature, and among the three precipitation "
        "variables, means that the marginal associations reported here cannot be "
        "translated into independent coefficient estimates in a multivariate regression "
        "without first pruning predictors to a near-orthogonal set.  Third, omitted "
        "variables such as public holidays, promotional events, and local footfall "
        "patterns may confound the weather-revenue relationship; the correlations "
        "documented here motivate including weather as a predictor in a revenue model "
        "but do not constitute causal evidence."
    )

    # ── Save and report ────────────────────────────────────────────────────────
    doc.save(REPORT)

    paras_after = len(doc.paragraphs)
    images_after = len(doc.inline_shapes)
    print(f"Paragraphs: {paras_before} -> {paras_after}  (+{paras_after - paras_before})")
    print(f"Images:     {images_before} -> {images_after}  (+{images_after - images_before})")

    all_h1 = [(i, p.text) for i, p in enumerate(doc.paragraphs) if p.style.name == "Heading 1"]
    print("\nAll H1 headings:")
    for i, t in all_h1:
        print(f"  [{i:3d}] {t}")

    all_h2 = [(i, p.text) for i, p in enumerate(doc.paragraphs) if p.style.name == "Heading 2"]
    print("\nAll H2 headings:")
    for i, t in all_h2:
        print(f"  [{i:3d}] {t}")

    print("\nFirst 3 paras of new section:")
    for p in doc.paragraphs[paras_before + 1 : paras_before + 5]:
        if p.text.strip():
            print(f"  {p.text[:100]}")


if __name__ == "__main__":
    main()
